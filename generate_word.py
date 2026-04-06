#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公共建筑碳排放计算报告 - Word文档生成器
基于GB/T 51366-2019和JS/T 303-2026
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# 设置默认样式
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.space_after = Pt(6)
paragraph_format.line_spacing = 1.5

# ========================================
# 封面页
# ========================================
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('公共建筑碳排放计算报告')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run.bold = True
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于GB/T 51366-2019《建筑碳排放计算标准》\n基于JS/T 303-2026《公共机构碳排放核算指南》')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')

# 项目信息表
info_table = doc.add_table(rows=8, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ['项目名称：', '__________________________'],
    ['建筑地址：', '__________________________'],
    ['建筑类型：', '□办公 □商业 □医院 □学校 □酒店 □其他'],
    ['建筑面积：', '__________ m²'],
    ['计算年度：', '__________年'],
    ['编制单位：', '__________________________'],
    ['审 核 人：', '__________________________'],
    ['编制日期：', '______年____月____日'],
]

for i, (label, value) in enumerate(info_data):
    p0 = info_table.rows[i].cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run0 = p0.add_run(label)
    run0.font.size = Pt(14)
    run0.bold = True
    run0.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run1 = info_table.rows[i].cells[1].paragraphs[0].add_run(value)
    run1.font.size = Pt(14)

doc.add_page_break()

# ========================================
# 目录
# ========================================
doc.add_heading('目 录', level=1)
doc.add_paragraph('')

toc_items = [
    '1. 概述', '2. 核算依据与标准', '3. 项目基本信息', '4. 核算边界与范围',
    '5. 运行阶段碳排放计算', '   5.1 能源消耗汇总', '   5.2 暖通空调系统',
    '   5.3 生活热水系统', '   5.4 照明系统', '   5.5 电梯系统',
    '   5.6 可再生能源系统', '6. 建材生产及运输阶段碳排放计算',
    '7. 建造及拆除阶段碳排放计算', '8. 碳排放汇总与强度分析',
    '9. 减排措施与建议', '10. 核算结论', '附录A：计算参数表',
    '附录B：能源碳排放因子表', '附录C：建材碳排放因子表',
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    if not item.startswith('   '):
        p.runs[0].bold = True

doc.add_page_break()

# ========================================
# 正文内容
# ========================================
doc.add_heading('1. 概述', level=1)
doc.add_heading('1.1 编制目的', level=2)
doc.add_paragraph('为贯彻国家有关应对气候变化和节能减排的方针政策，规范公共建筑碳排放计算方法，节约资源，保护环境，特编制本碳排放计算报告。')

doc.add_heading('2. 核算依据与标准', level=1)
standards = [
    'GB/T 51366-2019《建筑碳排放计算标准》',
    'JS/T 303-2026《公共机构碳排放核算指南》',
    'GB/T 24040《环境管理 生命周期评价 原则与框架》',
    'GB/T 24044《环境管理 生命周期评价 要求与指南》',
    'GB 50189《公共建筑节能设计标准》',
]
for std in standards:
    doc.add_paragraph(std, style='List Bullet')

doc.add_heading('3. 项目基本信息', level=1)
basic_table = doc.add_table(rows=10, cols=4)
basic_table.style = 'Table Grid'
basic_table.alignment = WD_TABLE_ALIGNMENT.CENTER
basic_data = [
    ['项目名称', '', '建筑类型', ''],
    ['建筑地址', '', '结构类型', ''],
    ['建筑面积', 'm²', '建筑高度', 'm'],
    ['地上面积', 'm²', '地下面积', 'm²'],
    ['地上层数', '层', '地下层数', '层'],
    ['竣工年份', '年', '计算年度', '年'],
    ['使用人数', '人', '年运行天数', '天'],
    ['绿地面积', 'm²', '绿地率', '%'],
    ['年用电量', 'kWh', '年用气量', 'm³'],
    ['年用热量', 'GJ', '可再生能源系统', ''],
]

for i, row_data in enumerate(basic_data):
    for j, cell_data in enumerate(row_data):
        p = basic_table.rows[i].cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_data)
        run.font.size = Pt(11)
        if j % 2 == 0:
            run.bold = True

doc.add_page_break()

doc.add_heading('4. 核算边界与范围', level=1)
doc.add_paragraph('核算边界涵盖以下三个阶段：')
phases = [
    ('建材生产及运输阶段', '建筑主体结构材料、围护结构材料、构件和部品等'),
    ('建造及拆除阶段', '分部分项工程施工碳排放、措施项目碳排放、拆除碳排放'),
    ('运行阶段', '暖通空调、生活热水、照明及电梯、可再生能源系统碳排放'),
]
for title, desc in phases:
    p = doc.add_paragraph()
    run = p.add_run(f'（{title}）')
    run.bold = True
    p.add_run(desc)

doc.add_heading('5. 运行阶段碳排放计算', level=1)
doc.add_paragraph('根据GB/T 51366-2019第4章，运行阶段碳排放计算包括：')
doc.add_heading('5.1 能源消耗汇总', level=2)
energy_table = doc.add_table(rows=9, cols=6)
energy_table.style = 'Table Grid'
energy_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['能源类型', '年消耗量', '单位', '折标准煤(tce)', '碳排放因子', '碳排放量(tCO₂)']
for j, header in enumerate(headers):
    p = energy_table.rows[0].cells[j].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

energy_types = [
    ['电力', '________', 'kWh', '________', '0.5836 tCO₂/MWh', '________'],
    ['天然气', '________', 'm³', '________', '2.184 kgCO₂/m³', '________'],
    ['集中供热', '________', 'GJ', '________', '0.11 tCO₂/GJ', '________'],
    ['汽油', '________', 'L', '________', '2.179 kgCO₂/L', '________'],
    ['柴油', '________', 'L', '________', '2.718 kgCO₂/L', '________'],
    ['液化石油气', '________', 'kg', '________', '3.166 kgCO₂/kg', '________'],
    ['其他能源', '________', '', '________', '________', '________'],
    ['合计', '', '', '________', '', '________'],
]
for i, row_data in enumerate(energy_types, 1):
    for j, cell_data in enumerate(row_data):
        p = energy_table.rows[i].cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_data)
        run.font.size = Pt(10)
        if i == 8:
            run.bold = True

doc.add_heading('5.2 暖通空调系统', level=2)
doc.add_paragraph('制冷剂年碳排放量计算公式：Cr = mr / ye × GWPr / 1000')
doc.add_paragraph('制冷剂年碳排放量：________ tCO₂e')

doc.add_heading('5.3 生活热水系统', level=2)
doc.add_paragraph('年耗热量计算公式：Qr = 4.187 × m × qr × ρr × (tr - t1) × T / 1000')
doc.add_paragraph('生活热水系统年碳排放量：________ tCO₂e')

doc.add_heading('5.4 照明系统', level=2)
doc.add_paragraph('照明能耗计算公式：El = Σ(Pi,j × Ai × ti,j) / 1000 + 24 × Pp × A / 1000')
doc.add_paragraph('照明系统年碳排放量：________ tCO₂e')

doc.add_heading('5.5 电梯系统', level=2)
doc.add_paragraph('电梯能耗计算公式：Ee = (3.6 × P × ta × V × W + Estandby × ts) / 1000')
doc.add_paragraph('电梯系统年碳排放量：________ tCO₂e')

doc.add_heading('5.6 可再生能源系统', level=2)
doc.add_paragraph('光伏系统年发电量计算公式：Epv = I × KE × (1 - KS) × Ap')
doc.add_paragraph('光伏系统年减碳量：________ tCO₂e')

doc.add_page_break()

doc.add_heading('6. 建材生产及运输阶段', level=1)
doc.add_paragraph('建材生产碳排放计算公式：Csc = Σ(Mi × Fi)')
doc.add_paragraph('建材运输碳排放计算公式：Cys = Σ(Mi × Di × Ti)')
doc.add_paragraph('建材生产阶段碳排放总计：________ tCO₂e')
doc.add_paragraph('建材运输阶段碳排放总计：________ tCO₂e')

doc.add_heading('7. 建造及拆除阶段', level=1)
doc.add_paragraph('建造阶段碳排放计算公式：CJZ = Σ(Ejz,i × EFi) / A')
doc.add_paragraph('建造阶段碳排放总计：________ tCO₂e')
doc.add_paragraph('拆除阶段碳排放总计：________ tCO₂e')

doc.add_heading('8. 碳排放汇总与强度分析', level=1)
summary_table = doc.add_table(rows=11, cols=4)
summary_table.style = 'Table Grid'
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
s_headers = ['阶段', '碳排放量(tCO₂e)', '单位面积碳排放(kgCO₂e/m²)', '占比(%)']
for j, header in enumerate(s_headers):
    p = summary_table.rows[0].cells[j].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

summary_data = [
    ['1. 建材生产阶段', '________', '________', '________'],
    ['2. 建材运输阶段', '________', '________', '________'],
    ['3. 建造阶段', '________', '________', '________'],
    ['4. 拆除阶段', '________', '________', '________'],
    ['5. 运行阶段-电力', '________', '________', '________'],
    ['6. 运行阶段-化石燃料', '________', '________', '________'],
    ['7. 运行阶段-热力', '________', '________', '________'],
    ['8. 可再生能源减碳', '________', '________', '________'],
    ['建筑全生命周期碳排放总计', '________', '________', '100%'],
]
for i, row_data in enumerate(summary_data, 1):
    for j, cell_data in enumerate(row_data):
        p = summary_table.rows[i].cells[j].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_data)
        run.font.size = Pt(11)

doc.add_heading('9. 减排措施与建议', level=1)
doc.add_paragraph('1. 优化暖通空调系统运行策略，提高能效')
doc.add_paragraph('2. 增加可再生能源利用比例')
doc.add_paragraph('3. 加强照明系统节能改造')
doc.add_paragraph('4. 推进绿色建筑建材采购')

doc.add_heading('10. 核算结论', level=1)
doc.add_paragraph('建筑全生命周期碳排放总量为 ________ tCO₂e。')
doc.add_paragraph('单位面积碳排放强度为 ________ kgCO₂e/m²。')
doc.add_paragraph('碳排放强度等级评定为：□优秀 □良好 □合格 □不合格')

doc.add_page_break()
doc.add_heading('附录A：计算参数表', level=1)
doc.add_heading('附录B：能源碳排放因子表', level=1)
doc.add_heading('附录C：建材碳排放因子表', level=1)

# 签署页
doc.add_page_break()
doc.add_heading('报告签署页', level=1)
doc.add_paragraph('')
sign_table = doc.add_table(rows=5, cols=2)
sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sign_data = [
    ['编制人', '________________'],
    ['审核人', '________________'],
    ['批准人', '________________'],
    ['编制单位（盖章）', ''],
    ['日 期', '______年____月____日'],
]
for i, (label, value) in enumerate(sign_data):
    sign_table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
    sign_table.rows[i].cells[1].paragraphs[0].add_run(value)
    for cell in sign_table.rows[i].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存
word_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "公共建筑碳排放计算报告.docx")
doc.save(word_path)
print(f"✅ Word文档已保存至: {word_path}")
