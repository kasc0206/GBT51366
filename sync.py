#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
OpenXML Excel-Word 数据同步模块
基于 openpyxl 和 python-docx 实现跨平台数据同步

功能：
  - Excel → Word: 将Excel单元格数据同步到Word文档
  - Word → Excel: 将Word表单数据回写到Excel
  - 支持占位符替换、表格同步、图表更新

用法：
  python sync.py --excel data.xlsx --word template.docx --output output.docx
"""

import os
import re
import logging
import argparse
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass, field
from enum import Enum

import openpyxl
from openpyxl.utils import get_column_letter, column_index_from_string
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class SyncDirection(Enum):
    """同步方向"""
    EXCEL_TO_WORD = "excel_to_word"
    WORD_TO_EXCEL = "word_to_excel"


@dataclass
class CellMapping:
    """单元格映射配置"""
    excel_cell: str          # Excel单元格引用，如 "A1", "B5"
    word_placeholder: str    # Word占位符，如 "{{项目名称}}"
    sheet_name: str = ""     # Excel工作表名（可选）
    format_type: str = ""    # 格式化类型：text, number, date, currency
    sheet_index: int = 0     # 工作表索引


@dataclass
class TableMapping:
    """表格映射配置"""
    excel_range: str         # Excel数据范围，如 "A1:D10"
    word_table_index: int    # Word表格索引（从0开始）
    sheet_name: str = ""
    header_row: bool = True  # 是否包含表头


@dataclass
class SyncConfig:
    """同步配置"""
    mappings: List[CellMapping] = field(default_factory=list)
    table_mappings: List[TableMapping] = field(default_factory=list)
    default_sheet_name: str = ""
    placeholder_pattern: str = r"\{\{(.+?)\}\}"  # 占位符正则：{{key}}


class ExcelReader:
    """Excel 读取器"""
    
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.wb = openpyxl.load_workbook(filepath, data_only=True)
        logger.info(f"已加载 Excel 文件: {filepath}")
    
    def get_cell_value(self, cell_ref: str, sheet_name: str = "") -> Any:
        """
        获取单元格值
        
        Args:
            cell_ref: 单元格引用，如 "A1", "B5"
            sheet_name: 工作表名
            
        Returns:
            单元格值
        """
        ws = self._get_worksheet(sheet_name)
        try:
            value = ws[cell_ref].value
            logger.debug(f"读取 {sheet_name or '默认'}!{cell_ref} = {value}")
            return value
        except Exception as e:
            logger.warning(f"读取单元格 {cell_ref} 失败: {e}")
            return None
    
    def get_range_values(self, range_str: str, sheet_name: str = "") -> List[List[Any]]:
        """
        获取范围值
        
        Args:
            range_str: 范围字符串，如 "A1:D10"
            sheet_name: 工作表名
            
        Returns:
            二维列表数据
        """
        ws = self._get_worksheet(sheet_name)
        try:
            data = []
            for row in ws[range_str]:
                row_data = [cell.value for cell in row]
                data.append(row_data)
            logger.debug(f"读取范围 {range_str}: {len(data)} 行")
            return data
        except Exception as e:
            logger.warning(f"读取范围 {range_str} 失败: {e}")
            return []
    
    def get_named_range(self, name: str) -> Any:
        """获取命名范围的值"""
        try:
            if name in self.wb.defined_names:
                ref = self.wb.defined_names[name].attr_text
                ws_name, cell_ref = ref.split('!')
                ws_name = ws_name.strip("'")
                return self.get_cell_value(cell_ref, ws_name)
        except Exception as e:
            logger.warning(f"读取命名范围 {name} 失败: {e}")
        return None
    
    def get_sheet_names(self) -> List[str]:
        """获取所有工作表名称"""
        return self.wb.sheetnames
    
    def _get_worksheet(self, sheet_name: str = ""):
        """获取工作表"""
        if sheet_name and sheet_name in self.wb.sheetnames:
            return self.wb[sheet_name]
        return self.wb.active


class WordProcessor:
    """Word 处理器"""
    
    def __init__(self, template_path: str = ""):
        self.doc = Document(template_path) if template_path else Document()
        if template_path:
            logger.info(f"已加载 Word 模板: {template_path}")
    
    def replace_placeholders(self, data: Dict[str, Any], 
                            pattern: str = r"\{\{(.+?)\}\}") -> int:
        """
        替换占位符
        
        Args:
            data: 数据字典 {占位符名称: 值}
            pattern: 占位符正则模式
            
        Returns:
            替换次数
        """
        replace_count = 0
        
        for paragraph in self.doc.paragraphs:
            for match in re.finditer(pattern, paragraph.text):
                key = match.group(1).strip()
                if key in data and data[key] is not None:
                    value = self._format_value(data[key])
                    self._replace_in_run(paragraph, match.group(0), str(value))
                    replace_count += 1
                    logger.debug(f"替换占位符: {{{{{key}}}}} -> {value}")
        
        # 处理表格中的占位符
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for match in re.finditer(pattern, paragraph.text):
                            key = match.group(1).strip()
                            if key in data and data[key] is not None:
                                value = self._format_value(data[key])
                                self._replace_in_run(paragraph, match.group(0), str(value))
                                replace_count += 1
        
        logger.info(f"共替换 {replace_count} 个占位符")
        return replace_count
    
    def update_table(self, table_index: int, data: List[List[Any]], 
                    header_row: bool = True):
        """
        更新Word表格
        
        Args:
            table_index: 表格索引
            data: 表格数据（二维列表）
            header_row: 是否将第一行作为表头
        """
        if table_index >= len(self.doc.tables):
            logger.warning(f"表格索引 {table_index} 超出范围")
            return
        
        table = self.doc.tables[table_index]
        
        # 清空现有表格
        self._clear_table(table)
        
        # 写入新数据
        for i, row_data in enumerate(data):
            if i >= len(table.rows):
                table.add_row()
            
            row = table.rows[i]
            for j, value in enumerate(row_data):
                if j < len(row.cells):
                    cell = row.cells[j]
                    cell.text = str(value) if value is not None else ""
                    
                    # 格式化表头
                    if header_row and i == 0:
                        self._format_header_cell(cell)
        
        logger.info(f"已更新表格 {table_index}: {len(data)} 行")
    
    def save(self, output_path: str):
        """保存文档"""
        self.doc.save(output_path)
        logger.info(f"已保存 Word 文件: {output_path}")
    
    def _replace_in_run(self, paragraph, old_text: str, new_text: str):
        """在段落中替换文本"""
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
    
    def _format_value(self, value: Any) -> str:
        """格式化值"""
        if isinstance(value, float):
            # 判断是否为百分比
            if value < 1:
                return f"{value*100:.1f}%"
            return f"{value:,.2f}"
        elif isinstance(value, int):
            return f"{value:,}"
        return str(value)
    
    def _clear_table(self, table):
        """清空表格内容"""
        for row in table.rows:
            for cell in row.cells:
                cell.text = ""
    
    def _format_header_cell(self, cell):
        """格式化表头单元格"""
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)


class SyncManager:
    """同步管理器"""
    
    def __init__(self, config: SyncConfig):
        self.config = config
        self.excel: Optional[ExcelReader] = None
        self.word: Optional[WordProcessor] = None
    
    def load_excel(self, filepath: str):
        """加载Excel文件"""
        self.excel = ExcelReader(filepath)
    
    def load_word(self, template_path: str = ""):
        """加载Word模板"""
        self.word = WordProcessor(template_path)
    
    def sync_excel_to_word(self, output_path: str) -> Dict:
        """
        Excel → Word 同步
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            同步统计信息
        """
        if not self.excel or not self.word:
            raise ValueError("请先加载Excel和Word文件")
        
        stats = {
            "cells_synced": 0,
            "tables_synced": 0,
            "errors": []
        }
        
        # 1. 同步单元格数据
        data_dict = {}
        for mapping in self.config.mappings:
            sheet = mapping.sheet_name or self.config.default_sheet_name
            value = self.excel.get_cell_value(mapping.excel_cell, sheet)
            data_dict[mapping.word_placeholder] = value
            stats["cells_synced"] += 1
        
        # 2. 替换Word占位符
        self.word.replace_placeholders(data_dict, self.config.placeholder_pattern)
        
        # 3. 同步表格数据
        for table_mapping in self.config.table_mappings:
            sheet = table_mapping.sheet_name or self.config.default_sheet_name
            data = self.excel.get_range_values(table_mapping.excel_range, sheet)
            self.word.update_table(
                table_mapping.word_table_index, 
                data,
                table_mapping.header_row
            )
            stats["tables_synced"] += 1
        
        # 4. 保存
        self.word.save(output_path)
        
        logger.info(f"同步完成: {stats['cells_synced']} 个单元格, {stats['tables_synced']} 个表格")
        return stats
    
    def sync_word_to_excel(self, output_path: str) -> Dict:
        """
        Word → Excel 同步（从Word表单读取数据写入Excel）
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            同步统计信息
        """
        # TODO: 实现反向同步
        logger.warning("Word → Excel 反向同步尚未实现")
        return {"cells_synced": 0, "errors": ["反向同步未实现"]}


def create_default_config() -> SyncConfig:
    """创建默认同步配置"""
    config = SyncConfig(default_sheet_name="项目基本信息")
    
    # 添加默认映射
    config.mappings = [
        CellMapping("D4", "{{项目名称}}"),
        CellMapping("D5", "{{建筑地址}}"),
        CellMapping("D6", "{{建筑类型}}"),
        CellMapping("D7", "{{建筑面积}}"),
        CellMapping("D14", "{{计算年度}}"),
        CellMapping("D18", "{{年用电量}}"),
        CellMapping("D19", "{{年用气量}}"),
        CellMapping("D20", "{{年用热量}}"),
    ]
    
    return config


def main():
    """CLI 入口"""
    parser = argparse.ArgumentParser(description="OpenXML Excel-Word 数据同步工具")
    parser.add_argument("--excel", "-e", required=True, help="Excel文件路径")
    parser.add_argument("--word", "-w", required=True, help="Word模板路径")
    parser.add_argument("--output", "-o", required=True, help="输出文件路径")
    parser.add_argument("--config", "-c", help="配置文件路径（JSON）")
    parser.add_argument("--direction", "-d", default="excel_to_word",
                       choices=["excel_to_word", "word_to_excel"],
                       help="同步方向")
    parser.add_argument("--verbose", "-v", action="store_true", help="详细输出")
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    # 创建配置
    config = create_default_config()
    
    # 创建同步管理器
    sync = SyncManager(config)
    sync.load_excel(args.excel)
    sync.load_word(args.word)
    
    # 执行同步
    if args.direction == "excel_to_word":
        stats = sync.sync_excel_to_word(args.output)
        print(f"\n✅ 同步完成!")
        print(f"   单元格同步: {stats['cells_synced']} 个")
        print(f"   表格同步: {stats['tables_synced']} 个")
        print(f"   输出文件: {args.output}")
    else:
        stats = sync.sync_word_to_excel(args.output)
        print(f"\n⚠️  Word → Excel 同步尚未完全实现")


if __name__ == "__main__":
    main()
