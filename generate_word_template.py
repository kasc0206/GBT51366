#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
公共建筑碳排放计算报告 - Word模板生成器（适配精细化版）
生成与精细化Excel计算表格配套的Word报告模板

用法：
    python generate_word_template.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "templates", "公共建筑碳排放计算报告.docx")


def add_styled_table(doc, headers, rows, col_widths=None):
    """添加样式表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(10)
        run.font.name = '微软雅黑'
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(
            '<w:shd {} w:fill="1F4E79"/>'.format(nsdecls('w'))
        )
        cell._tc.get_or_add_tcPr().append(shading)

    # 数据行
    for i, row_data in enumerate(rows, 1):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(cell_data))
            run.font.size = Pt(10)
            run.font.name = '微软雅黑'

    # 列宽
    if col_widths:
        for row in table.rows:
            for j, width in enumerate(col_widths):
                row.cells[j].width = Cm(width)

    doc.add_paragraph('')  # 空行分隔


def main():
    """生成Word模板"""
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # ========================================
    # 封面
    # ========================================
    for _ in range(6):
        doc.add_paragraph('')

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('公共建筑碳排放计算报告')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('基于 GB/T 51366-2019《建筑碳排放计算标准》\n基于 JS/T 303-2026《公共机构碳排放核算指南》')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')

    # 项目信息表
    info_table = doc.add_table(rows=8, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ['项目名称：', '{{项目名称}}'],
        ['建筑地址：', '{{建筑地址}}'],
        ['建筑类型：', '{{建筑类型}}'],
        ['建筑面积：', '{{建筑面积}} m²'],
        ['计算年度：', '{{计算年度}}年'],
        ['编制单位：', '__________________________'],
        ['审 核 人：', '__________________________'],
        ['编制日期：', '______年____月____日'],
    ]

    for i, (label, value) in enumerate(info_data):
        p0 = info_table.rows[i].cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run0 = p0.add_run(label)
        run0.font.size = Pt(13)
        run0.bold = True
        run0.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        p1 = info_table.rows[i].cells[1].paragraphs[0]
        run1 = p1.add_run(value)
        run1.font.size = Pt(13)

    for cell in info_table.rows[0].cells + info_table.rows[1].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            '<w:tcBorders %s>'
            '<w:top w:val="none" w:sz="0" w:space="0"/>'
            '<w:left w:val="none" w:sz="0" w:space="0"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0"/>'
            '<w:right w:val="none" w:sz="0" w:space="0"/>'
            '</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(tcBorders)

    doc.add_page_break()

    # ========================================
    # 目录
    # ========================================
    doc.add_heading('目 录', level=1)
    doc.add_paragraph('')

    toc_items = [
        '1. 概述',
        '2. 核算依据与标准',
        '3. 项目基本信息',
        '4. 核算边界与范围',
        '5. 运行阶段碳排放计算',
        '   5.1 负荷计算',
        '   5.2 暖通空调系统',
        '   5.3 生活热水系统',
        '   5.4 照明系统',
        '   5.5 电梯及其他设备',
        '   5.6 可再生能源系统',
        '6. 建材生产及运输阶段碳排放计算',
        '   6.1 建材生产阶段',
        '   6.2 建材运输阶段',
        '7. 建造及拆除阶段碳排放计算',
        '   7.1 土方及结构工程',
        '   7.2 装饰装修及机电安装',
        '   7.3 拆除阶段',
        '8. 水资源与废弃物碳排放',
        '9. 绿地碳汇与隐含碳排放',
        '10. 碳排放汇总与强度分析',
        '11. 减排措施与建议',
        '12. 核算结论',
        '附录A：计算参数表',
        '附录B：能源碳排放因子表',
        '附录C：建材碳排放因子表',
    ]

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        if not item.startswith('   '):
            p.runs[0].bold = True

    doc.add_page_break()

    # ========================================
    # 1. 概述
    # ========================================
    doc.add_heading('1. 概述', level=1)

    doc.add_heading('1.1 编制目的', level=2)
    doc.add_paragraph(
        '为贯彻国家有关应对气候变化和节能减排的方针政策，规范公共建筑碳排放计算方法，'
        '节约资源，保护环境，特编制本碳排放计算报告。本报告采用精细化计算单元方式，'
        '将建筑全生命周期碳排放拆分为39个独立计算单元，实现单一职责、独立完整、'
        '易于扩展和维护的计算体系。'
    )

    doc.add_heading('1.2 适用范围', level=2)
    doc.add_paragraph(
        '本报告适用于新建、扩建和改建的民用建筑的运行、建造及拆除、建材生产及运输阶段的'
        '碳排放计算。本报告针对公共建筑类型，包括办公建筑、商业建筑、医院、学校、酒店等。'
    )

    # ========================================
    # 2. 核算依据与标准
    # ========================================
    doc.add_heading('2. 核算依据与标准', level=1)

    doc.add_heading('2.1 主要标准', level=2)
    standards = [
        'GB/T 51366-2019《建筑碳排放计算标准》',
        'JS/T 303-2026《公共机构碳排放核算指南》',
        'GB/T 24040《环境管理 生命周期评价 原则与框架》',
        'GB/T 24044《环境管理 生命周期评价 要求与指南》',
        'GB 50189《公共建筑节能设计标准》',
        'GB 50555《民用建筑节水设计标准》',
        'GB 50034《建筑照明设计标准》',
    ]
    for std in standards:
        doc.add_paragraph(std, style='List Bullet')

    # ========================================
    # 3. 项目基本信息
    # ========================================
    doc.add_heading('3. 项目基本信息', level=1)

    doc.add_heading('3.1 建筑基本信息', level=2)
    basic_headers = ['项目', '数值', '单位', '数据来源']
    basic_rows = [
        ['项目名称', '{{项目名称}}', '-', '项目文件'],
        ['建筑地址', '{{建筑地址}}', '-', '规划许可证'],
        ['建筑类型', '{{建筑类型}}', '-', '设计文件'],
        ['建筑面积', '{{建筑面积}}', 'm²', '房产证'],
        ['使用人数', '{{使用人数}}', '人', '物业记录'],
        ['年运行天数', 365, '天', '实际运行'],
        ['绿地面积', '{{绿地面积}}', 'm²', '规划许可证'],
        ['光伏装机容量', '{{光伏装机容量}}', 'kWp', '设计文件'],
    ]
    add_styled_table(doc, basic_headers, basic_rows, [15, 12, 8, 12])

    doc.add_heading('3.2 气象参数', level=2)
    weather_headers = ['参数', '数值', '单位', '数据来源']
    weather_rows = [
        ['年太阳辐射照度', '{{年太阳辐射照度}}', 'kWh/m²', '气象局'],
        ['采暖室外计算温度', '{{采暖室外计算温度}}', '°C', '设计规范'],
        ['空调室外计算温度', '{{空调室外计算温度}}', '°C', '设计规范'],
        ['冬季室外平均风速', '{{冬季室外平均风速}}', 'm/s', '气象局'],
    ]
    add_styled_table(doc, weather_headers, weather_rows, [15, 12, 8, 12])

    doc.add_page_break()

    # ========================================
    # 4. 核算边界与范围
    # ========================================
    doc.add_heading('4. 核算边界与范围', level=1)

    doc.add_heading('4.1 核算边界', level=2)
    doc.add_paragraph(
        '根据GB/T 51366-2019第3.0.1条，本核算以单栋建筑（或建筑群）为计算对象。'
        '核算边界涵盖以下三个阶段：'
    )

    phases = [
        ('建材生产及运输阶段', '建筑主体结构材料、围护结构材料、构件和部品等主要建材的生产过程碳排放和运输过程碳排放。'),
        ('建造及拆除阶段', '分部分项工程施工碳排放、措施项目碳排放、拆除碳排放。'),
        ('运行阶段', '暖通空调、生活热水、照明及电梯、可再生能源系统碳排放。'),
    ]
    for title, desc in phases:
        p = doc.add_paragraph()
        run = p.add_run(f'（{title}）')
        run.bold = True
        p.add_run(desc)

    # ========================================
    # 5. 运行阶段碳排放计算
    # ========================================
    doc.add_heading('5. 运行阶段碳排放计算', level=1)

    doc.add_heading('5.1 负荷计算', level=2)

    doc.add_heading('5.1.1 冷负荷计算', level=3)
    doc.add_paragraph('年冷负荷计算公式：年冷负荷 = 建筑面积 × 冷负荷指标 × 年供冷小时数 × 同时使用系数 / 1000')

    cooling_headers = ['参数', '数值', '单位']
    cooling_rows = [
        ['建筑面积', '{{建筑面积}}', 'm²'],
        ['冷负荷指标', '{{冷负荷指标}}', 'W/m²'],
        ['年供冷小时数', '{{年供冷小时数}}', 'h'],
        ['同时使用系数', 0.8, '-'],
        ['年冷负荷', '{{年冷负荷}}', 'kWh'],
    ]
    add_styled_table(doc, cooling_headers, cooling_rows, [15, 12, 8])

    doc.add_heading('5.1.2 热负荷计算', level=3)
    doc.add_paragraph('年热负荷计算公式：年热负荷 = 建筑面积 × 热负荷指标 × 年供暖小时数 × 同时使用系数 / 1000')

    heating_headers = ['参数', '数值', '单位']
    heating_rows = [
        ['建筑面积', '{{建筑面积}}', 'm²'],
        ['热负荷指标', '{{热负荷指标}}', 'W/m²'],
        ['年供暖小时数', '{{年供暖小时数}}', 'h'],
        ['同时使用系数', 0.75, '-'],
        ['年热负荷', '{{年热负荷}}', 'kWh'],
    ]
    add_styled_table(doc, heating_headers, heating_rows, [15, 12, 8])

    doc.add_heading('5.2 暖通空调系统', level=2)

    doc.add_heading('5.2.1 冷源系统', level=3)
    doc.add_paragraph('冷源系统能耗 = 年冷负荷 / COP + 输配系统能耗')

    hvac_headers = ['参数', '数值', '单位']
    hvac_rows = [
        ['年冷负荷', '{{年冷负荷}}', 'kWh'],
        ['冷源系统COP', '{{冷源COP}}', '-'],
        ['冷冻水泵功率', '{{冷冻水泵功率}}', 'kW'],
        ['冷却水泵功率', '{{冷却水泵功率}}', 'kW'],
        ['冷却塔功率', '{{冷却塔功率}}', 'kW'],
        ['冷源系统总能耗', '{{冷源系统能耗}}', 'kWh'],
    ]
    add_styled_table(doc, hvac_headers, hvac_rows, [15, 12, 8])

    doc.add_heading('5.2.2 制冷剂排放', level=3)
    doc.add_paragraph('制冷剂年排放量 = 充注量 × 台数 / 设备寿命 × 年泄漏率')

    ref_headers = ['参数', '数值', '单位']
    ref_rows = [
        ['制冷剂类型', '{{制冷剂类型}}', '-'],
        ['GWP值', '{{GWP值}}', '-'],
        ['充注量', '{{制冷剂充注量}}', 'kg/台'],
        ['设备台数', '{{设备台数}}', '台'],
        ['设备寿命', 15, '年'],
        ['年泄漏率', 5, '%'],
        ['制冷剂年排放量', '{{制冷剂年排放}}', 'tCO₂e'],
    ]
    add_styled_table(doc, ref_headers, ref_rows, [15, 12, 8])

    doc.add_heading('5.3 生活热水系统', level=2)
    doc.add_paragraph('年耗热量 = 4.187 × 用水人数 × 用水定额 × (设计热水温度 - 设计冷水温度) / 24 × 年使用天数')

    hw_headers = ['参数', '数值', '单位']
    hw_rows = [
        ['用水人数', '{{用水人数}}', '人'],
        ['用水定额', 50, 'L/人·天'],
        ['设计热水温度', 60, '°C'],
        ['设计冷水温度', 10, '°C'],
        ['年使用天数', 365, '天'],
        ['热源效率', 0.9, '-'],
        ['生活热水年能耗', '{{生活热水年能耗}}', 'kWh'],
    ]
    add_styled_table(doc, hw_headers, hw_rows, [15, 12, 8])

    doc.add_heading('5.4 照明系统', level=2)
    doc.add_paragraph('照明年能耗 = Σ(面积 × 功率密度 × 日运行时间 × 年运行天数) / 1000')

    light_headers = ['区域', '面积(m²)', '功率密度(W/m²)', '日运行时间(h)', '年运行天数']
    light_rows = [
        ['办公室', '{{办公室面积}}', 9, 10, 250],
        ['会议室', '{{会议室面积}}', 11, 8, 250],
        ['大堂门厅', '{{大堂面积}}', 15, 14, 365],
        ['走廊', '{{走廊面积}}', 5, 14, 365],
        ['地下车库', '{{车库面积}}', 2, 24, 365],
        ['照明年能耗合计', '{{照明年能耗}}', 'kWh', '', ''],
    ]
    add_styled_table(doc, light_headers, light_rows, [12, 10, 12, 10, 10])

    doc.add_heading('5.5 电梯及其他设备', level=2)

    doc.add_heading('5.5.1 电梯系统', level=3)
    elevator_headers = ['参数', '数值', '单位']
    elevator_rows = [
        ['客梯数量', '{{客梯数量}}', '台'],
        ['客梯载重', '{{客梯载重}}', 'kg'],
        ['客梯速度', '{{客梯速度}}', 'm/s'],
        ['货梯数量', '{{货梯数量}}', '台'],
        ['特定能量消耗', '{{电梯特定能耗}}', 'mWh/kgm'],
        ['年运行时间', '{{电梯年运行时间}}', 'h'],
        ['电梯年能耗', '{{电梯年能耗}}', 'kWh'],
    ]
    add_styled_table(doc, elevator_headers, elevator_rows, [15, 12, 8])

    doc.add_heading('5.5.2 数据中心', level=3)
    dc_headers = ['参数', '数值', '单位']
    dc_rows = [
        ['IT设备总功率', '{{IT设备功率}}', 'kW'],
        ['PUE值', 1.8, '-'],
        ['年运行时间', 8760, 'h'],
        ['数据中心年能耗', '{{数据中心年能耗}}', 'kWh'],
    ]
    add_styled_table(doc, dc_headers, dc_rows, [15, 12, 8])

    doc.add_heading('5.6 可再生能源系统', level=2)

    doc.add_heading('5.6.1 光伏发电', level=3)
    doc.add_paragraph('年发电量 = 光伏面板面积 × 年辐射照度 × 转换效率 × (1 - 系统损失率)')

    pv_headers = ['参数', '数值', '单位']
    pv_rows = [
        ['光伏面板面积', '{{光伏面积}}', 'm²'],
        ['年辐射照度', '{{年太阳辐射照度}}', 'kWh/m²'],
        ['转换效率', 18, '%'],
        ['系统损失率', 20, '%'],
        ['年发电量', '{{光伏年发电量}}', 'kWh'],
        ['电网排放因子', 0.5836, 'kgCO₂/kWh'],
        ['年减碳量', '{{光伏年减碳量}}', 'tCO₂'],
    ]
    add_styled_table(doc, pv_headers, pv_rows, [15, 12, 8])

    doc.add_page_break()

    # ========================================
    # 6. 建材生产及运输阶段
    # ========================================
    doc.add_heading('6. 建材生产及运输阶段碳排放计算', level=1)

    doc.add_heading('6.1 建材生产阶段', level=2)
    doc.add_paragraph('建材生产碳排放 = Σ(消耗量 × 排放因子 × (1 - 回收率))')

    mat_headers = ['建材类别', '消耗量', '单位', '排放因子', '回收率(%)', '碳排放(tCO₂e)']
    mat_rows = [
        ['钢筋', '{{钢筋消耗量}}', 't', '2.033 tCO₂e/t', 90, '{{钢筋碳排放}}'],
        ['混凝土', '{{混凝土消耗量}}', 'm³', '0.320 tCO₂e/m³', 30, '{{混凝土碳排放}}'],
        ['围护建材', '{{围护建材消耗量}}', '按需', '按类型', '按需', '{{围护建材碳排放}}'],
        ['装饰材料', '{{装饰材料消耗量}}', '按需', '按类型', '按需', '{{装饰材料碳排放}}'],
        ['安装材料', '{{安装材料消耗量}}', '按需', '按类型', '按需', '{{安装材料碳排放}}'],
        ['建材生产合计', '', '', '', '', '{{建材生产总碳排放}}'],
    ]
    add_styled_table(doc, mat_headers, mat_rows, [12, 10, 8, 12, 8, 12])

    doc.add_heading('6.2 建材运输阶段', level=2)
    doc.add_paragraph('建材运输碳排放 = Σ(运输量 × 运输距离 × 运输排放因子)')

    trans_headers = ['建材', '运输量(t)', '距离(km)', '运输方式', '排放因子', '碳排放(tCO₂e)']
    trans_rows = [
        ['钢筋', '{{钢筋运输量}}', '{{钢筋运输距离}}', '公路重载', 0.0697, '{{钢筋运输碳排放}}'],
        ['混凝土', '{{混凝土运输量}}', '{{混凝土运输距离}}', '商砼罐车', 0.0712, '{{混凝土运输碳排放}}'],
        ['其他材料', '{{其他材料运输量}}', '{{其他材料运输距离}}', '公路综合', 0.0667, '{{其他运输碳排放}}'],
        ['建材运输合计', '', '', '', '', '{{建材运输总碳排放}}'],
    ]
    add_styled_table(doc, trans_headers, trans_rows, [10, 10, 10, 10, 10, 12])

    doc.add_page_break()

    # ========================================
    # 7. 建造及拆除阶段
    # ========================================
    doc.add_heading('7. 建造及拆除阶段碳排放计算', level=1)

    doc.add_heading('7.1 土方及结构工程', level=2)

    const_headers = ['工程类别', '工程量', '单位', '能耗系数(kWh/单位)', '碳排放(tCO₂e)']
    const_rows = [
        ['土方开挖', '{{土方开挖量}}', 'm³', 5.5, '{{土方开挖碳排放}}'],
        ['土方回填', '{{土方回填量}}', 'm³', 3.2, '{{土方回填碳排放}}'],
        ['基础混凝土', '{{基础混凝土量}}', 'm³', 35, '{{基础混凝土碳排放}}'],
        ['柱梁板混凝土', '{{柱梁板混凝土量}}', 'm³', 40, '{{柱梁板碳排放}}'],
        ['钢筋制作安装', '{{结构钢筋量}}', 't', 120, '{{结构钢筋碳排放}}'],
        ['土方及结构工程合计', '', '', '', '{{土方结构总碳排放}}'],
    ]
    add_styled_table(doc, const_headers, const_rows, [15, 12, 8, 15, 12])

    doc.add_heading('7.2 装饰装修及机电安装', level=2)

    deco_headers = ['工程类别', '工程量', '单位', '能耗系数(kWh/单位)', '碳排放(tCO₂e)']
    deco_rows = [
        ['砌体工程', '{{砌体工程量}}', 'm³', 15, '{{砌体碳排放}}'],
        ['外墙保温', '{{外墙保温面积}}', 'm²', 2.5, '{{外墙保温碳排放}}'],
        ['内墙装饰', '{{内墙装饰面积}}', 'm²', 1.8, '{{内墙碳排放}}'],
        ['楼地面工程', '{{楼地面面积}}', 'm²', 2.2, '{{楼地面碳排放}}'],
        ['给排水系统', '{{给排水项数}}', '项', 5000, '{{给排水碳排放}}'],
        ['电气系统', '{{电气项数}}', '项', 8000, '{{电气碳排放}}'],
        ['暖通系统', '{{暖通项数}}', '项', 12000, '{{暖通碳排放}}'],
        ['装饰装修及机电合计', '', '', '', '{{装饰机电总碳排放}}'],
    ]
    add_styled_table(doc, deco_headers, deco_rows, [15, 12, 8, 15, 12])

    doc.add_heading('7.3 拆除阶段', level=2)

    demo_headers = ['拆除项目', '工程量', '单位', '能耗系数(kWh/单位)', '碳排放(tCO₂e)']
    demo_rows = [
        ['人工拆除', '{{人工拆除面积}}', 'm²', 8, '{{人工拆除碳排放}}'],
        ['机械拆除', '{{机械拆除面积}}', 'm²', 25, '{{机械拆除碳排放}}'],
        ['垃圾外运', '{{垃圾外运量}}', 't', 15, '{{垃圾外运碳排放}}'],
        ['拆除工程合计', '', '', '', '{{拆除总碳排放}}'],
    ]
    add_styled_table(doc, demo_headers, demo_rows, [15, 12, 8, 15, 12])

    doc.add_page_break()

    # ========================================
    # 8. 水资源与废弃物
    # ========================================
    doc.add_heading('8. 水资源与废弃物碳排放', level=1)

    doc.add_heading('8.1 水资源消耗', level=2)

    water_headers = ['用水类型', '年用水量(m³)', '供水能耗(kWh/m³)', '碳排放(tCO₂)']
    water_rows = [
        ['生活用水', '{{生活用水量}}', 0.35, '{{生活用水碳排放}}'],
        ['空调补水', '{{空调补水量}}', 0.35, '{{空调补水碳排放}}'],
        ['绿化灌溉', '{{绿化灌溉量}}', 0.35, '{{绿化灌溉碳排放}}'],
        ['水资源合计', '', '', '{{水资源总碳排放}}'],
    ]
    add_styled_table(doc, water_headers, water_rows, [12, 12, 12, 12])

    doc.add_heading('8.2 废弃物处理', level=2)

    waste_headers = ['废弃物类型', '年产生量(t)', '处理排放因子(kgCO₂/t)', '碳排放(tCO₂)']
    waste_rows = [
        ['生活垃圾填埋', '{{生活垃圾量}}', 350, '{{生活垃圾碳排放}}'],
        ['餐厨垃圾', '{{餐厨垃圾量}}', -80, '{{餐厨垃圾碳排放}}'],
        ['可回收物', '{{可回收物量}}', -800, '{{可回收物碳排放}}'],
        ['建筑垃圾', '{{建筑垃圾量}}', 15, '{{建筑垃圾碳排放}}'],
        ['废弃物合计', '', '', '{{废弃物总碳排放}}'],
    ]
    add_styled_table(doc, waste_headers, waste_rows, [12, 12, 15, 12])

    doc.add_page_break()

    # ========================================
    # 9. 绿地碳汇与隐含碳
    # ========================================
    doc.add_heading('9. 绿地碳汇与隐含碳排放', level=1)

    doc.add_heading('9.1 绿地碳汇', level=2)
    doc.add_paragraph('建筑年碳汇 = Σ(绿化面积 × 碳汇系数)')

    sink_headers = ['碳汇类型', '面积(m²)', '碳汇系数[kgCO₂/(m²·a)]', '年碳汇量(tCO₂)']
    sink_rows = [
        ['乔木绿地', '{{乔木绿地面积}}', 2.5, '{{乔木碳汇}}'],
        ['灌木绿地', '{{灌木绿地面积}}', 1.2, '{{灌木碳汇}}'],
        ['草坪', '{{草坪面积}}', 0.5, '{{草坪碳汇}}'],
        ['屋顶绿化', '{{屋顶绿化面积}}', 1.8, '{{屋顶碳汇}}'],
        ['垂直绿化', '{{垂直绿化面积}}', 1.5, '{{垂直碳汇}}'],
        ['建筑年碳汇总计', '', '', '{{绿地碳汇总计}}'],
    ]
    add_styled_table(doc, sink_headers, sink_rows, [12, 12, 18, 12])

    doc.add_heading('9.2 隐含碳排放', level=2)
    doc.add_paragraph('隐含碳年化 = Σ(设备数量 × 碳排放因子 / 使用寿命)')

    emb_headers = ['资产类别', '数量', '单位', '碳排放因子', '使用寿命(年)', '年碳排放(tCO₂e)']
    emb_rows = [
        ['办公桌', '{{办公桌数量}}', '套', '85 kgCO₂e/套', 10, '{{办公桌隐含碳}}'],
        ['电脑', '{{电脑数量}}', '台', '250 kgCO₂e/台', 5, '{{电脑隐含碳}}'],
        ['空调', '{{空调数量}}', '台', '450 kgCO₂e/台', 10, '{{空调隐含碳}}'],
        ['隐含碳合计', '', '', '', '', '{{隐含碳年排放}}'],
    ]
    add_styled_table(doc, emb_headers, emb_rows, [12, 8, 8, 12, 10, 12])

    doc.add_page_break()

    # ========================================
    # 10. 碳排放汇总与强度分析
    # ========================================
    doc.add_heading('10. 碳排放汇总与强度分析', level=1)

    doc.add_heading('10.1 全生命周期碳排放汇总', level=2)

    sum_headers = ['排放阶段', '碳排放量(tCO₂e)', '单位面积(kgCO₂e/m²)', '占比(%)', '来源工作表']
    sum_rows = [
        ['建材生产阶段', '{{建材生产碳排放}}', '{{建材生产单位面积}}', '{{建材生产占比}}', '24-28表'],
        ['建材运输阶段', '{{建材运输碳排放}}', '{{建材运输单位面积}}', '{{建材运输占比}}', '29表'],
        ['土方及结构工程', '{{土方结构碳排放}}', '{{土方结构单位面积}}', '{{土方结构占比}}', '30-31表'],
        ['装饰装修及机电', '{{装饰机电碳排放}}', '{{装饰机电单位面积}}', '{{装饰机电占比}}', '32-33表'],
        ['拆除工程', '{{拆除碳排放}}', '{{拆除单位面积}}', '{{拆除占比}}', '34表'],
        ['运行阶段', '{{运行阶段碳排放}}', '{{运行单位面积}}', '{{运行占比}}', '05-22表'],
        ['水资源与废弃物', '{{水废弃物碳排放}}', '{{水废弃物单位面积}}', '{{水废弃物占比}}', '35表'],
        ['隐含碳排放', '{{隐含碳排放}}', '{{隐含碳单位面积}}', '{{隐含碳占比}}', '37表'],
        ['减：绿地碳汇', '{{绿地碳汇}}', '{{碳汇单位面积}}', '{{碳汇占比}}', '36表(负值)'],
        ['建筑全生命周期总计', '{{全生命周期总碳排放}}', '{{全生命周期单位面积}}', '100%', '38表'],
    ]
    add_styled_table(doc, sum_headers, sum_rows, [15, 12, 12, 8, 10])

    doc.add_heading('10.2 碳排放强度指标', level=2)

    intensity_headers = ['指标名称', '计算值', '单位', '国家标准限额', '达标情况']
    intensity_rows = [
        ['单位面积建材生产碳排放', '{{建材生产强度}}', 'kgCO₂e/m²', '<200', '{{建材生产达标}}'],
        ['单位面积运行阶段碳排放', '{{运行阶段强度}}', 'kgCO₂e/(m²·a)', '<75', '{{运行阶段达标}}'],
        ['单位面积全生命周期碳排放', '{{全生命周期强度}}', 'kgCO₂e/m²', '<850', '{{全生命周期达标}}'],
        ['人均运行阶段碳排放', '{{人均碳排放}}', 'tCO₂/(人·a)', '<5.0', '{{人均达标}}'],
        ['碳排放强度等级', '{{碳排放等级}}', '级', '-', '□优秀 □良好 □合格 □不合格'],
    ]
    add_styled_table(doc, intensity_headers, intensity_rows, [20, 10, 12, 10, 15])

    doc.add_page_break()

    # ========================================
    # 11. 减排措施与建议
    # ========================================
    doc.add_heading('11. 减排措施与建议', level=1)

    measures = [
        ('暖通空调系统优化', [
            '提高冷热源设备能效，选用能效等级1级设备',
            '优化系统运行策略，避免能源浪费',
            '加强建筑围护结构保温，降低冷热负荷',
            '推广使用变频技术和智能控制系统',
        ]),
        ('生活热水系统节能', [
            '优先采用太阳能热水系统',
            '提高热水管道保温性能，降低热损失',
            '推广使用节水器具，减少热水消耗',
        ]),
        ('照明系统节能', [
            '全面采用LED等高效光源',
            '推广智能照明控制系统，实现分区、分时控制',
            '充分利用自然采光，减少人工照明时间',
        ]),
        ('可再生能源利用', [
            '扩大光伏发电系统装机容量',
            '推广地源热泵系统应用',
            '提高可再生能源在总能耗中的占比',
        ]),
        ('绿色建材采购', [
            '优先选用绿色建材认证产品',
            '提高再生材料使用比例',
            '减少建材运输距离，优先就地取材',
        ]),
    ]

    for title, items in measures:
        p = doc.add_paragraph()
        run = p.add_run(f'（{title}）')
        run.bold = True
        for item in items:
            doc.add_paragraph(item, style='List Bullet 2')

    doc.add_page_break()

    # ========================================
    # 12. 核算结论
    # ========================================
    doc.add_heading('12. 核算结论', level=1)

    doc.add_heading('12.1 核算结果总结', level=2)
    doc.add_paragraph('经核算，本建筑全生命周期碳排放情况如下：')

    conclusions = [
        f'建筑全生命周期碳排放总量为 {{全生命周期总碳排放}} tCO₂e。',
        f'其中，运行阶段碳排放 {{运行阶段碳排放}} tCO₂e，占比 {{运行占比}}%。',
        f'建材生产及运输阶段碳排放 {{建材生产运输碳排放}} tCO₂e，占比 {{建材生产运输占比}}%。',
        f'建造及拆除阶段碳排放 {{建造拆除碳排放}} tCO₂e，占比 {{建造拆除占比}}%。',
        f'单位面积碳排放强度为 {{全生命周期强度}} kgCO₂e/m²。',
        f'人均碳排放强度为 {{人均碳排放}} tCO₂e/(人·a)。',
        f'可再生能源减碳量为 {{可再生能源减碳量}} tCO₂e，占总排放的 {{可再生能源占比}}%。',
    ]

    for conclusion in conclusions:
        doc.add_paragraph(conclusion)

    doc.add_heading('12.2 碳排放等级评定', level=2)
    doc.add_paragraph(
        '根据核算结果，本建筑碳排放强度等级评定为：{{碳排放等级}}'
    )

    doc.add_paragraph(
        '□ 优秀：单位面积碳排放低于国家先进标准限额\n'
        '□ 良好：单位面积碳排放低于国家现行标准限额\n'
        '□ 合格：单位面积碳排放满足国家现行标准限额\n'
        '□ 不合格：单位面积碳排放超过国家现行标准限额'
    )

    doc.add_heading('12.3 持续改进建议', level=2)
    doc.add_paragraph(
        '为持续降低建筑碳排放，建议采取以下措施：'
    )

    improvements = [
        '建立碳排放监测平台，实现能耗数据实时监测和分析',
        '定期开展能源审计，识别节能潜力点',
        '制定碳排放削减计划，明确目标和时间表',
        '加强建筑使用人员节能意识培训',
        '推进碳排放权交易，利用市场机制降低碳排放',
        '探索碳中和路径，通过购买碳抵消实现建筑碳中和',
    ]

    for i, item in enumerate(improvements, 1):
        doc.add_paragraph(f'{i}. {item}')

    doc.add_page_break()

    # ========================================
    # 附录
    # ========================================
    doc.add_heading('附录A：计算参数表', level=1)

    doc.add_heading('A.1 室内设计参数', level=2)
    param_headers = ['参数', '数值', '单位', '标准依据']
    param_rows = [
        ['夏季室内设计温度', 26, '°C', 'GB 50189'],
        ['冬季室内设计温度', 20, '°C', 'GB 50189'],
        ['夏季室内相对湿度', 65, '%', 'GB 50189'],
        ['人均新风量', 30, 'm³/(h·人)', 'GB 50189'],
        ['办公室照明标准', 500, 'lux', 'GB 50034'],
    ]
    add_styled_table(doc, param_headers, param_rows, [18, 10, 12, 12])

    doc.add_heading('附录B：能源碳排放因子表', level=1)

    ef_headers = ['能源类型', '单位', '碳排放因子', '数据来源']
    ef_rows = [
        ['电力', 'kWh', '0.5836 kgCO₂/kWh', '区域电网'],
        ['天然气', 'm³', '2.184 kgCO₂/m³', 'GB/T 51366-2019'],
        ['集中供热', 'GJ', '0.11 tCO₂/GJ', 'JS/T 303-2026'],
        ['汽油', 'L', '2.179 kgCO₂/L', 'JS/T 303-2026'],
        ['柴油', 'L', '2.718 kgCO₂/L', 'JS/T 303-2026'],
    ]
    add_styled_table(doc, ef_headers, ef_rows, [15, 8, 18, 15])

    doc.add_heading('附录C：建材碳排放因子表', level=1)

    mf_headers = ['建材名称', '单位', '碳排放因子', '回收率']
    mf_rows = [
        ['钢筋(HRB400)', 't', '2.033 tCO₂e/t', '90%'],
        ['混凝土(C30)', 'm³', '0.320 tCO₂e/m³', '30%'],
        ['加气混凝土砌块', 'm³', '0.225 tCO₂e/m³', '30%'],
        ['XPS保温板', 'm³', '0.095 tCO₂e/m³', '0%'],
        ['玻璃（Low-E中空）', 'm²', '0.015 tCO₂e/m²', '80%'],
        ['铝合金门窗', 'm²', '0.025 tCO₂e/m²', '90%'],
        ['花岗岩', 'm²', '0.038 tCO₂e/m²', '50%'],
        ['瓷砖', 'm²', '0.012 tCO₂e/m²', '30%'],
        ['电缆（铜芯）', 'm', '0.002 tCO₂e/m', '95%'],
    ]
    add_styled_table(doc, mf_headers, mf_rows, [18, 8, 18, 10])

    doc.add_paragraph('')
    doc.add_paragraph('注：以上碳排放因子数据为缺省值，如有经第三方审核的建材碳足迹数据，应优先采用实测数据。').italic = True

    # ========================================
    # 签署页
    # ========================================
    doc.add_page_break()
    doc.add_heading('报告签署页', level=1)

    doc.add_paragraph('')
    doc.add_paragraph('')

    sign_table = doc.add_table(rows=5, cols=2)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sign_data = [
        ['编制人', '________________'],
        ['审核人', '________________'],
        ['批准人', '________________'],
        ['编制单位（盖章）', ''],
        ['日 期', '______年____月____日'],
    ]

    for i, (label, value) in enumerate(sign_data):
        sign_table.rows[i].cells[0].paragraphs[0].add_run(label).bold = True
        sign_table.rows[i].cells[1].paragraphs[0].add_run(value)
        for cell in sign_table.rows[i].cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 保存
    doc.save(OUTPUT_FILE)
    print(f"✅ Word模板已保存至: {OUTPUT_FILE}")
    print("📝 模板中的 {{占位符}} 可通过 sync.py 模块与Excel数据同步")


if __name__ == "__main__":
    main()
